
# report_service\app.pyfrom flask import Flask, request, send_file, jsonify
from fpdf import FPDF
import tempfile

app = Flask(__name__)

@app.route('/report/export', methods=['POST'])
def export():
    data = request.get_json()
    patient_name = data.get('patient_name')
    history = data.get('history')
    format_type = data.get('format')  # 'pdf', 'docx', 'print'

    if not patient_name or not history:
        return jsonify({'error': 'Datos insuficientes'}), 400

    if format_type == 'pdf':
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt=f"Historial Médico de {patient_name}", ln=True, align='C')
        pdf.multi_cell(0, 10, txt=history)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp:
            pdf.output(tmp.name)
            return send_file(tmp.name, as_attachment=True, attachment_filename=f"Historial_{patient_name}.pdf")

    elif format_type == 'docx':
        from docx import Document
        document = Document()
        document.add_heading(f'Historial Médico de {patient_name}', 0)
        document.add_paragraph(history)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            document.save(tmp.name)
            return send_file(tmp.name, as_attachment=True, attachment_filename=f"Historial_{patient_name}.docx")

    elif format_type == 'print':
        # Lógica para imprimir directamente
        return jsonify({'message': 'Enviado a imprimir'}), 200
    else:
        return jsonify({'error': 'Formato no soportado'}), 400

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5004)
